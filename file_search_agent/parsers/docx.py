"""
DOCX Parser - Extracts text and metadata from Word documents.
Uses python-docx for parsing.
"""

from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseParser, ParsedDocument, DocumentChunk


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""
    
    supported_extensions = [".docx", ".doc"]
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a Word document and extract text with structure awareness."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() == ".doc":
            # .doc files need special handling (legacy format)
            return self._parse_legacy_doc(file_path)
        
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted DOCX file: {file_path}")
        
        # Extract core properties (metadata)
        core_props = doc.core_properties
        
        # Extract text with structure
        content_parts = []
        section_data = []
        current_section = "Introduction"
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            if para.style and para.style.name.startswith("Heading"):
                current_section = text
            
            content_parts.append(text)
            section_data.append({
                "text": text,
                "section": current_section,
                "style": para.style.name if para.style else "Normal"
            })
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content_parts.append(row_text)
                    section_data.append({
                        "text": row_text,
                        "section": current_section,
                        "style": "Table"
                    })
        
        full_content = "\n\n".join(content_parts)
        
        # Create section-aware chunks
        chunks = self._create_section_aware_chunks(section_data)
        
        # Get file metadata
        file_meta = self.get_file_metadata(file_path)
        
        return ParsedDocument(
            file_path=file_path,
            file_name=file_path.name,
            file_type="docx",
            content=full_content,
            chunks=chunks,
            file_size=file_meta["file_size"],
            created_at=core_props.created,
            modified_at=core_props.modified,
            title=core_props.title or file_path.stem,
            author=core_props.author,
            page_count=None,  # DOCX doesn't have fixed pages
            metadata={
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
            }
        )
    
    def _create_section_aware_chunks(
        self,
        section_data: List[Dict[str, Any]]
    ) -> List[DocumentChunk]:
        """Create chunks that respect document sections."""
        chunks = []
        chunk_index = 0
        current_chunk_text = ""
        current_section = ""
        char_position = 0
        chunk_start = 0
        
        for item in section_data:
            text = item["text"]
            section = item["section"]
            
            # If section changes, start a new chunk
            if section != current_section and current_chunk_text:
                chunks.append(DocumentChunk(
                    text=current_chunk_text.strip(),
                    chunk_index=chunk_index,
                    start_char=chunk_start,
                    end_char=char_position,
                    section=current_section,
                    metadata={"section": current_section}
                ))
                chunk_index += 1
                current_chunk_text = ""
                chunk_start = char_position
            
            current_section = section
            
            # Check if adding this text exceeds chunk size
            if len(current_chunk_text) + len(text) > self.chunk_size:
                if current_chunk_text:
                    chunks.append(DocumentChunk(
                        text=current_chunk_text.strip(),
                        chunk_index=chunk_index,
                        start_char=chunk_start,
                        end_char=char_position,
                        section=current_section,
                        metadata={"section": current_section}
                    ))
                    chunk_index += 1
                chunk_start = char_position
                current_chunk_text = text + "\n\n"
            else:
                current_chunk_text += text + "\n\n"
            
            char_position += len(text) + 2
        
        # Add final chunk
        if current_chunk_text.strip():
            chunks.append(DocumentChunk(
                text=current_chunk_text.strip(),
                chunk_index=chunk_index,
                start_char=chunk_start,
                end_char=char_position,
                section=current_section,
                metadata={"section": current_section}
            ))
        
        return chunks
    
    def _parse_legacy_doc(self, file_path: Path) -> ParsedDocument:
        """
        Parse legacy .doc files.
        Requires antiword or similar tool installed on the system.
        """
        import subprocess
        
        try:
            # Try using antiword
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                check=True
            )
            content = result.stdout
        except (subprocess.CalledProcessError, FileNotFoundError):
            # Fallback: try reading as plain text (usually doesn't work well)
            raise ValueError(
                f"Cannot parse legacy .doc file: {file_path}. "
                "Install 'antiword' or convert to .docx format."
            )
        
        file_meta = self.get_file_metadata(file_path)
        chunks = self.chunk_text(content)
        
        return ParsedDocument(
            file_path=file_path,
            file_name=file_path.name,
            file_type="doc",
            content=content,
            chunks=chunks,
            file_size=file_meta["file_size"],
            created_at=file_meta.get("created_at"),
            modified_at=file_meta.get("modified_at"),
            title=file_path.stem,
            author=None,
        )
